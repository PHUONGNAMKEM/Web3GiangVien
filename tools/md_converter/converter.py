"""
MD Converter - Chuyen doi Markdown -> PDF / DOCX
Converter module: su dung fpdf2 (Unicode TTF) + python-docx
"""

import os
import re
import tempfile
import urllib.request
import urllib.parse
from PIL import Image as PILImage
import markdown
from io import BytesIO
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ============================================================
# MARKDOWN TO HTML (for web preview)
# ============================================================

def md_to_html(md_content: str) -> str:
    """Convert Markdown to styled HTML (for preview only)"""
    extensions = [
        'markdown.extensions.tables',
        'markdown.extensions.fenced_code',
        'markdown.extensions.codehilite',
        'markdown.extensions.toc',
        'markdown.extensions.nl2br',
        'markdown.extensions.sane_lists',
        'markdown.extensions.smarty',
        'markdown.extensions.attr_list',
    ]
    extension_configs = {
        'markdown.extensions.codehilite': {
            'css_class': 'highlight',
            'guess_lang': False,
        },
    }
    html_body = markdown.markdown(
        md_content,
        extensions=extensions,
        extension_configs=extension_configs,
        output_format='html5'
    )
    return html_body


# ============================================================
# PDF CONVERSION (fpdf2 - native Unicode TTF support)
# ============================================================

class VietnamesePDF(FPDF):
    """Custom PDF class with Vietnamese font support and styled rendering."""

    def __init__(self, metadata=None):
        self.metadata = metadata or {}
        style = self.metadata.get('style', {})
        orientation = style.get('orientation', 'P')
        page_size = style.get('page_size', 'A4')
        
        super().__init__(orientation=orientation, format=page_size)
        
        # Thiết lập margin
        margin_map = {'narrow': 15, 'standard': 20, 'wide': 25}
        self.margin_val = margin_map.get(style.get('margin', 'standard'), 20)
        self.set_margins(self.margin_val, self.margin_val, self.margin_val)
        self.set_auto_page_break(auto=True, margin=self.margin_val)
        
        # Thiết lập màu sắc động
        self._setup_colors()
        self._setup_fonts()

    def _setup_colors(self):
        style = self.metadata.get('style', {})
        theme_hex = style.get('theme_color', '#0f3460')
        custom_enabled = style.get('custom_colors_enabled', False)
        
        def hex_to_rgb(hex_str):
            if not hex_str:
                return (15, 52, 96)
            hex_str = hex_str.lstrip('#')
            if len(hex_str) == 3:
                hex_str = ''.join(c*2 for c in hex_str)
            return tuple(int(hex_str[i:i+2], 16) for i in (0, 2, 4))
            
        try:
            self.color_primary = hex_to_rgb(theme_hex)
        except Exception:
            self.color_primary = (15, 52, 96) # fallback
            
        if custom_enabled:
            try:
                self.color_heading1 = hex_to_rgb(style.get('h1_color'))
                self.color_heading2 = hex_to_rgb(style.get('h2_color'))
                self.color_heading3 = hex_to_rgb(style.get('h3_color'))
                self.color_primary = hex_to_rgb(style.get('table_header_bg'))
                self.color_border = hex_to_rgb(style.get('hr_color'))
                self.color_bold = hex_to_rgb(style.get('bold_color'))
                self.color_toc = hex_to_rgb(style.get('toc_color'))
                self.color_quote_border = hex_to_rgb(style.get('quote_border'))
                self.color_quote_bg = hex_to_rgb(style.get('quote_bg'))
            except Exception:
                # Fallback to automatic if error parsing
                r, g, b = self.color_primary
                self.color_heading1 = self.color_primary
                self.color_heading2 = (max(0, r - 30), max(0, g - 30), max(0, b - 30))
                self.color_heading3 = (max(0, r - 50), max(0, g - 50), min(255, b + 40))
                self.color_border = (max(0, r - 10), min(255, g + 20), min(255, b + 20))
                self.color_bold = self.color_heading2
                self.color_toc = self.color_heading1
                self.color_quote_bg = (248, 246, 255)
                self.color_quote_border = self.color_heading3
        else:
            r, g, b = self.color_primary
            self.color_heading1 = self.color_primary
            
            if theme_hex.lower() == '#222222':
                # B&W (grayscale) theme
                self.color_heading1 = (0, 0, 0)
                self.color_heading2 = (34, 34, 34)
                self.color_heading3 = (68, 68, 68)
                self.color_border = (153, 153, 153)
                self.color_bold = (0, 0, 0)
                self.color_toc = (0, 0, 0)
                self.color_quote_bg = (245, 245, 245)
                self.color_quote_border = (85, 85, 85)
                self.color_primary = (34, 34, 34)
                self.color_code_text = (17, 17, 17)
                self.color_code_bg = (238, 238, 238)
                self.color_codeblock_bg = (240, 240, 240)
                self.color_codeblock_text = (17, 17, 17)
                self.color_link = (34, 34, 34)
            elif theme_hex.lower() == '#a91d22':
                self.color_heading2 = (169, 29, 34)
                self.color_heading3 = (245, 166, 35) # #f5a623
                self.color_border = (245, 166, 35)   # #f5a623
                self.color_bold = (169, 29, 34)
                self.color_toc = (169, 29, 34)
                self.color_quote_bg = (254, 251, 236) # #fefbec
                self.color_quote_border = (245, 166, 35)
            else:
                self.color_heading2 = (max(0, r - 30), max(0, g - 30), max(0, b - 30))
                self.color_heading3 = (max(0, r - 50), max(0, g - 50), min(255, b + 40))
                self.color_border = (max(0, r - 10), min(255, g + 20), min(255, b + 20))
                self.color_bold = self.color_heading2
                self.color_toc = self.color_heading1
                self.color_quote_bg = (248, 246, 255)
                self.color_quote_border = self.color_heading3
            
        self.color_text = (26, 26, 46)
        self.color_muted = (100, 100, 120)
        
        # Only set these if not already set by B&W theme
        if not hasattr(self, 'color_code_text') or theme_hex.lower() != '#222222':
            self.color_code_text = (233, 69, 96)
            self.color_code_bg = (240, 240, 245)
            self.color_codeblock_bg = (26, 26, 46)
            self.color_codeblock_text = (224, 224, 224)
            self.color_link = self.color_primary;

    def _setup_fonts(self):
        """Register Windows system fonts that support Vietnamese dynamically."""
        font_dir = os.path.join(os.environ.get('WINDIR', r'C:\Windows'), 'Fonts')
        
        style = self.metadata.get('style', {})
        font_family = style.get('font_family', 'Arial')
        
        font_map = {
            'Times New Roman': {
                'R': 'times.ttf',
                'B': 'timesbd.ttf',
                'I': 'timesi.ttf',
                'BI': 'timesbi.ttf'
            },
            'Calibri': {
                'R': 'calibri.ttf',
                'B': 'calibrib.ttf',
                'I': 'calibrii.ttf',
                'BI': 'calibriz.ttf'
            },
            'Courier New': {
                'R': 'cour.ttf',
                'B': 'courbd.ttf',
                'I': 'couri.ttf',
                'BI': 'courbi.ttf'
            },
            'Arial': {
                'R': 'arial.ttf',
                'B': 'arialbd.ttf',
                'I': 'ariali.ttf',
                'BI': 'arialbi.ttf'
            },
            'Segoe UI': {
                'R': 'segoeui.ttf',
                'B': 'segoeuib.ttf',
                'I': 'segoeuii.ttf',
                'BI': 'segoeuiz.ttf'
            },
            'Georgia': {
                'R': 'georgia.ttf',
                'B': 'georgiab.ttf',
                'I': 'georgiai.ttf',
                'BI': 'georgiaz.ttf'
            }
        }
        
        selected_font = font_map.get(font_family, font_map['Arial'])
        
        r_path = os.path.join(font_dir, selected_font['R'])
        b_path = os.path.join(font_dir, selected_font['B'])
        i_path = os.path.join(font_dir, selected_font['I'])
        bi_path = os.path.join(font_dir, selected_font['BI'])
        
        if not os.path.exists(r_path):
            fallback = font_map['Arial']
            r_path = os.path.join(font_dir, fallback['R'])
            b_path = os.path.join(font_dir, fallback['B'])
            i_path = os.path.join(font_dir, fallback['I'])
            bi_path = os.path.join(font_dir, fallback['BI'])
            
        self.add_font('VN', '', r_path)
        self.add_font('VN', 'B', b_path if os.path.exists(b_path) else r_path)
        self.add_font('VN', 'I', i_path if os.path.exists(i_path) else r_path)
        self.add_font('VN', 'BI', bi_path if os.path.exists(bi_path) else r_path)
        
        consolas = os.path.join(font_dir, 'consola.ttf')
        consolas_b = os.path.join(font_dir, 'consolab.ttf')
        if os.path.exists(consolas):
            self.add_font('Mono', '', consolas)
            self.add_font('Mono', 'B', consolas_b if os.path.exists(consolas_b) else consolas)
        else:
            self.add_font('Mono', '', r_path)

    def header(self):
        project_name = self.metadata.get('project_name', '')
        version = self.metadata.get('doc_version', '')
        logo_path = self.metadata.get('logo_path', '')  # Logo bên phải
        logo_left_path = self.metadata.get('logo_left_path', '')  # Logo bên trái
        
        if self.page_no() == 1:
            return
            
        pdf_w = self.w - self.l_margin - self.r_margin
        
        # Vẽ logo bên trái nếu có, khống chế chiều cao h=8mm để không đè lên dòng kẻ ngang y=20
        logo_left_w = 0
        if logo_left_path and os.path.exists(logo_left_path):
            try:
                w_calc_l = 12  # default fallback width
                with PILImage.open(logo_left_path) as img:
                    img_w, img_h = img.size
                    w_calc_l = (img_w / img_h) * 8  # scale to height 8mm
                if w_calc_l > 15:
                    w_calc_l = 15  # cap max width
                self.image(logo_left_path, x=self.l_margin, y=10, h=8, w=w_calc_l)
                logo_left_w = w_calc_l + 3  # width plus 3mm padding
            except Exception:
                pass
                
        # Vẽ logo bên phải nếu có, khống chế chiều cao h=8mm
        logo_right_w = 0
        if logo_path and os.path.exists(logo_path):
            try:
                w_calc_r = 12  # default fallback width
                with PILImage.open(logo_path) as img:
                    img_w, img_h = img.size
                    w_calc_r = (img_w / img_h) * 8  # scale to height 8mm
                if w_calc_r > 15:
                    w_calc_r = 15  # cap max width
                self.image(logo_path, x=self.w - self.r_margin - w_calc_r, y=10, h=8, w=w_calc_r)
                logo_right_w = w_calc_r + 3  # width plus 3mm padding
            except Exception:
                pass
                
        self.set_font('VN', 'I', 8)
        self.set_text_color(120, 120, 120)
        header_text = project_name
        if version:
            header_text += f" - Phiên bản: {version}"
            
            
        if header_text:
            self.set_x(self.l_margin + logo_left_w)
            cell_w = pdf_w - logo_left_w - logo_right_w
            self.cell(cell_w, 8, header_text, ln=0, align='L')
            
        # Separator line between header and content
        self.set_draw_color(180, 180, 180)
        self.set_line_width(0.3)
        self.line(self.l_margin, 20, self.w - self.r_margin, 20)
        self.ln(12)

    def footer(self):
        self.set_y(-15)
        self.set_font('VN', 'I', 8)
        self.set_text_color(150, 150, 150)

        # Separator line between content and footer
        self.set_draw_color(180, 180, 180)
        self.set_line_width(0.3)
        self.line(self.l_margin, self.get_y() - 2, self.w - self.r_margin, self.get_y() - 2)
        
        style = self.metadata.get('style', {})
        footer_format = style.get('footer_format', 'page_of_total')
        
        if footer_format == 'page_only':
            footer_text = f'{self.page_no()}'
            align = 'C'
        elif footer_format == 'brackets':
            footer_text = f'- {self.page_no()} -'
            align = 'C'
        elif footer_format == 'right_align':
            footer_text = f'Trang {self.page_no()}/{{nb}}'
            align = 'R'
        else: # page_of_total
            footer_text = f'Trang {self.page_no()}/{{nb}}'
            align = 'C'
            
        self.cell(0, 10, footer_text, align=align)


# --- Color palette ---
COLOR_HEADING1 = (15, 52, 96)       # #0f3460
COLOR_HEADING2 = (22, 33, 62)       # #16213e
COLOR_HEADING3 = (83, 52, 131)      # #533483
COLOR_HEADING4 = (15, 52, 96)       # #0f3460
COLOR_TEXT = (26, 26, 46)           # #1a1a2e
COLOR_MUTED = (100, 100, 120)
COLOR_CODE_TEXT = (233, 69, 96)     # #e94560
COLOR_CODE_BG = (240, 240, 245)     # #f0f0f5
COLOR_CODEBLOCK_BG = (26, 26, 46)   # #1a1a2e
COLOR_CODEBLOCK_TEXT = (224, 224, 224)
COLOR_LINK = (15, 52, 96)
COLOR_TABLE_HEADER = (15, 52, 96)   # #0f3460
COLOR_TABLE_HEADER_TEXT = (255, 255, 255)
COLOR_TABLE_ALT = (244, 246, 250)   # #f4f6fa
COLOR_BORDER = (233, 69, 96)        # #e94560
COLOR_QUOTE_BORDER = (83, 52, 131)  # #533483
COLOR_QUOTE_BG = (248, 246, 255)    # #f8f6ff


# ============================================================
# IMAGE AND TOC RENDERING HELPERS
# ============================================================

def _get_image_file(img_path: str) -> str:
    """Download image if URL, or find in ObsidianVault if local filename, returns local file path or None"""
    if img_path.startswith(('http://', 'https://')):
        try:
            temp_dir = tempfile.gettempdir()
            ext = os.path.splitext(urllib.parse.urlparse(img_path).path)[1]
            if not ext or len(ext) > 5:
                ext = '.png'
            temp_file = os.path.join(temp_dir, f"md_conv_tmp_{hash(img_path)}{ext}")
            
            req = urllib.request.Request(
                img_path, 
                headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
            )
            with urllib.request.urlopen(req, timeout=10) as response, open(temp_file, 'wb') as out_file:
                out_file.write(response.read())
            return temp_file
        except Exception as e:
            print(f"Error downloading image {img_path}: {e}")
            return None
    else:
        if os.path.exists(img_path):
            return img_path
        
        filename = os.path.basename(img_path)
        # Find workspace root
        root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        vault_dir = os.path.join(root_dir, "ObsidianVault")
        
        # Search in ObsidianVault first
        if os.path.exists(vault_dir):
            for r, dirs, files in os.walk(vault_dir):
                if filename in files:
                    return os.path.join(r, filename)
                
        # Search workspace
        for r, dirs, files in os.walk(root_dir):
            if any(p in r for p in ['.git', '.venv', '__pycache__', 'node_modules']):
                continue
            if filename in files:
                return os.path.join(r, filename)
                
        return None


def _pdf_image(pdf: FPDF, img_path: str, alt_text: str):
    """Render image into PDF using fpdf2, scaling to fit page width safely"""
    local_path = _get_image_file(img_path)
    if not local_path or not os.path.exists(local_path):
        pdf.ln(2)
        pdf.set_text_color(200, 50, 50)
        pdf.set_font('VN', 'I', 10)
        pdf.cell(0, 10, f"[Hình ảnh không tải được: {alt_text or img_path}]", ln=1)
        pdf.set_text_color(*COLOR_TEXT)
        pdf.ln(2)
        return
        
    try:
        with PILImage.open(local_path) as img:
            width_px, height_px = img.size
            
        page_width = pdf.w - pdf.l_margin - pdf.r_margin
        img_w = page_width
        img_h = (height_px / width_px) * img_w
        
        if pdf.get_y() + img_h > pdf.h - pdf.b_margin:
            pdf.add_page()
            
        pdf.ln(4)
        pdf.image(local_path, x=pdf.get_x(), y=pdf.get_y(), w=img_w)
        pdf.set_y(pdf.get_y() + img_h)
        pdf.ln(4)
    except Exception as e:
        print(f"Error rendering image in PDF: {e}")
        pdf.ln(2)
        pdf.set_text_color(200, 50, 50)
        pdf.set_font('VN', 'I', 10)
        pdf.cell(0, 10, f"[Lỗi hiển thị hình ảnh: {alt_text or img_path}]", ln=1)
        pdf.set_text_color(*COLOR_TEXT)
        pdf.ln(2)


def _pdf_toc(pdf: FPDF, headings_log: list, toc_pages: int):
    """Render the Table of Contents page(s) at the beginning of the PDF"""
    pdf.set_font('VN', 'B', 18)
    toc_color = getattr(pdf, 'color_toc', getattr(pdf, 'color_heading1', (15, 52, 96)))
    pdf.set_text_color(*toc_color)
    pdf.cell(0, 15, "Mục Lục", ln=1, align='C')
    pdf.ln(5)

    pdf.ln(8)
    
    for h in headings_log:
        level = h['level']
        if level > 3:
            continue
            
        text = h['text']
        page = h['page'] + toc_pages
        
        indent = (level - 1) * 8
        pdf.set_x(pdf.l_margin + indent)
        
        pdf.set_font('VN', 'B' if level == 1 else '', 11 if level == 1 else 10)
        
        # Chọn màu động dựa trên level đề mục hoặc gán chung theo màu TOC tùy chỉnh
        if hasattr(pdf, 'color_toc'):
            h_color = pdf.color_toc
        else:
            h_color = getattr(pdf, 'color_heading1', (15, 52, 96))
            if level == 2:
                h_color = getattr(pdf, 'color_heading2', (22, 33, 62))
            elif level == 3:
                h_color = getattr(pdf, 'color_heading3', (83, 52, 131))
            
        pdf.set_text_color(*h_color)
        
        line_w = pdf.w - pdf.l_margin - pdf.r_margin - indent
        text_w = pdf.get_string_width(text)
        page_str = str(page)
        page_w = pdf.get_string_width(page_str)
        
        dots_w = line_w - text_w - page_w - 6
        if dots_w > 0:
            dots_count = int(dots_w / pdf.get_string_width('.'))
            dots = '.' * dots_count
            pdf.write(6, text)
            pdf.write(6, f" {dots} ")
            pdf.write(6, page_str)
        else:
            pdf.write(6, text)
            pdf.set_x(pdf.w - pdf.r_margin - page_w)
            pdf.write(6, page_str)
            
        pdf.ln(8) # Tăng khoảng cách dòng tránh dính chữ
        
    pdf.add_page()


def convert_md_to_pdf(md_content: str, title: str = "Document", metadata: dict = None) -> bytes:
    """Convert Markdown content to PDF bytes with Vietnamese support via fpdf2."""
    metadata = metadata or {}
    
    # --- Pass 1: Render draft to count heading page numbers ---
    pdf_draft = VietnamesePDF(metadata)
    pdf_draft.alias_nb_pages()
    pdf_draft.add_page()
    pdf_draft.set_font('VN', '', 11)
    
    headings_log = []
    original_pdf_heading = _pdf_heading
    
    def log_pdf_heading(pdf, text, level):
        original_pdf_heading(pdf, text, level)
        headings_log.append({'text': text, 'level': level, 'page': pdf.page_no()})
        
    globals()['_pdf_heading'] = lambda pdf, text, level: log_pdf_heading(pdf, text, level)
    
    try:
        _run_pdf_rendering(pdf_draft, md_content)
    finally:
        globals()['_pdf_heading'] = original_pdf_heading
        
    # --- Pass 2: Final Render ---
    pdf = VietnamesePDF(metadata)
    pdf.alias_nb_pages()
    pdf.add_page()
    
    has_headings = any(h['level'] <= 3 for h in headings_log)
    if has_headings:
        toc_count = sum(1 for h in headings_log if h['level'] <= 3)
        toc_pages = (toc_count + 34) // 35
        _pdf_toc(pdf, headings_log, toc_pages)
        
    pdf.set_font('VN', '', 11)
    _run_pdf_rendering(pdf, md_content)
    
    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer.read()


def _run_pdf_rendering(pdf: FPDF, md_content: str):
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ''
    in_table = False
    table_header = []
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # --- Code block ---
        if line.strip().startswith('```'):
            if in_code_block:
                _pdf_code_block(pdf, '\n'.join(code_lines))
                code_lines = []
                code_lang = ''
                in_code_block = False
            else:
                if in_table and table_header:
                    _pdf_table(pdf, table_header, table_rows)
                    table_header, table_rows = [], []
                    in_table = False
                in_code_block = True
                lang_match = re.match(r'^```(\w+)', line.strip())
                code_lang = lang_match.group(1) if lang_match else ''
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- Table ---
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if all(re.match(r'^[-:]+$', c) for c in cells if c):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_header = cells
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table and table_header:
                _pdf_table(pdf, table_header, table_rows)
                table_header, table_rows = [], []
                in_table = False

        stripped = line.strip()

        # Empty line
        if not stripped:
            pdf.ln(2)
            i += 1
            continue

        # Horizontal rule — spacing only, no line
        if re.match(r'^[-*_]{3,}\s*$', stripped):
            pdf.ln(6)
            i += 1
            continue

        # Image block
        image_match = re.match(r'^!\[(.*?)\]\((.*?)\)', stripped)
        if image_match:
            alt = image_match.group(1)
            img_path = image_match.group(2)
            _pdf_image(pdf, img_path, alt)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            _pdf_heading(pdf, text, level)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            quote_text = re.sub(r'^>\s*', '', stripped)
            quote_text = re.sub(r'\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', r'[\1]', quote_text)
            _pdf_blockquote(pdf, quote_text)
            i += 1
            continue

        # List items
        list_match = re.match(r'^(\s*)([-*+]|\d+[.)])\s+(.*)', stripped)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            marker = list_match.group(2)
            text = list_match.group(3)
            is_ordered = bool(re.match(r'\d+[.)]', marker))
            _pdf_list_item(pdf, text, indent_level, is_ordered, marker)
            i += 1
            continue

        # Regular paragraph
        _pdf_paragraph(pdf, stripped)
        i += 1

    # Flush remaining
    if in_table and table_header:
        _pdf_table(pdf, table_header, table_rows)
    if in_code_block and code_lines:
        _pdf_code_block(pdf, '\n'.join(code_lines))


# ============================================================
# PDF RENDERING HELPERS
# ============================================================

def _pdf_heading(pdf: FPDF, text: str, level: int):
    """Render a heading."""
    sizes = {1: 22, 2: 16, 3: 13, 4: 11.5, 5: 11, 6: 10.5}
    colors = {
        1: getattr(pdf, 'color_heading1', (15, 52, 96)),
        2: getattr(pdf, 'color_heading2', (22, 33, 62)),
        3: getattr(pdf, 'color_heading3', (83, 52, 131)),
        4: getattr(pdf, 'color_primary', (15, 52, 96)),
        5: getattr(pdf, 'color_heading1', (15, 52, 96)),
        6: getattr(pdf, 'color_heading2', (22, 33, 62))
    }

    pdf.ln(4 if level <= 2 else 3)
    pdf.set_font('VN', 'B', sizes.get(level, 11))
    pdf.set_text_color(*colors.get(level, getattr(pdf, 'color_text', (26, 26, 46))))
    pdf.multi_cell(0, sizes.get(level, 11) * 0.5, text)

    # Spacing after heading (no underlines)
    if level <= 2:
        pdf.ln(2)
    else:
        pdf.ln(2)

    # Reset
    pdf.set_font('VN', '', 11)
    pdf.set_text_color(*getattr(pdf, 'color_text', (26, 26, 46)))


def _pdf_paragraph(pdf: FPDF, text: str):
    """Render a paragraph with inline formatting."""
    pdf.set_font('VN', '', 11)
    pdf.set_text_color(*getattr(pdf, 'color_text', (26, 26, 46)))
    _pdf_rich_text(pdf, text)
    pdf.ln(3)


def _pdf_rich_text(pdf: FPDF, text: str):
    """Render text with bold, italic, code, link formatting."""
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'
        r'|(\*\*(.+?)\*\*)'
        r'|(\*(.+?)\*)'
        r'|(`(.+?)`)'
        r'|(\[(.+?)\]\((.+?)\))'
    )

    last_end = 0
    line_h = 5.5

    for match in pattern.finditer(text):
        if match.start() > last_end:
            pdf.set_font('VN', '', 11)
            pdf.set_text_color(*getattr(pdf, 'color_text', (26, 26, 46)))
            pdf.write(line_h, text[last_end:match.start()])

        if match.group(2):  # bold+italic
            pdf.set_font('VN', 'BI', 11)
            pdf.set_text_color(*getattr(pdf, 'color_text', (26, 26, 46)))
            pdf.write(line_h, match.group(2))
        elif match.group(4):  # bold
            pdf.set_font('VN', 'B', 11)
            pdf.set_text_color(*getattr(pdf, 'color_bold', getattr(pdf, 'color_heading2', (22, 33, 62))))
            pdf.write(line_h, match.group(4))
        elif match.group(6):  # italic
            pdf.set_font('VN', 'I', 11)
            pdf.set_text_color(*getattr(pdf, 'color_heading3', (83, 52, 131)))
            pdf.write(line_h, match.group(6))
        elif match.group(8):  # inline code
            pdf.set_font('Mono', '', 9.5)
            pdf.set_text_color(*getattr(pdf, 'color_code_text', (233, 69, 96)))
            pdf.write(line_h, match.group(8))
        elif match.group(10):  # link
            pdf.set_font('VN', '', 11)
            pdf.set_text_color(*getattr(pdf, 'color_link', (15, 52, 96)))
            pdf.write(line_h, match.group(10))

        last_end = match.end()

    if last_end < len(text):
        pdf.set_font('VN', '', 11)
        pdf.set_text_color(*getattr(pdf, 'color_text', (26, 26, 46)))
        pdf.write(line_h, text[last_end:])

    pdf.ln(line_h)


def _pdf_code_block(pdf: FPDF, code_text: str):
    """Render a code block with dark background, handling page breaks.
    Two-pass per page chunk: draw background first, then text on top.
    """
    pdf.ln(3)
    x_start = pdf.l_margin
    width = pdf.w - pdf.l_margin - pdf.r_margin
    pdf.set_font('Mono', '', 8)
    
    code_lines = code_text.split('\n')
    line_h = 3.8
    padding_top = 4
    padding_bottom = 4
    max_y = pdf.h - pdf.b_margin - 18  # leave space for footer
    
    # Max chars per line
    char_w = pdf.get_string_width('M')
    max_chars = int((width - 14) / char_w) if char_w > 0 else 80
    
    # Truncate long lines
    rendered_lines = []
    for cl in code_lines:
        if len(cl) > max_chars:
            cl = cl[:max_chars - 3] + '...'
        rendered_lines.append(cl)
    
    # Check if current page has enough space for at least 4 lines
    # If not, move entirely to next page to avoid tiny chunks + big whitespace
    min_chunk_h = padding_top + 4 * line_h + padding_bottom
    remaining = max_y - pdf.get_y()
    if remaining < min_chunk_h and remaining < (padding_top + len(rendered_lines) * line_h + padding_bottom):
        pdf.add_page()
    
    # Split lines into page chunks
    chunks = []  # list of (list of lines)
    current_chunk = []
    y_cursor = pdf.get_y()
    
    for cl in rendered_lines:
        needed = line_h
        if len(current_chunk) == 0:
            needed += padding_top
        
        if y_cursor + needed + padding_bottom > max_y and len(current_chunk) > 0:
            chunks.append(current_chunk)
            current_chunk = []
            y_cursor = pdf.t_margin
        
        if len(current_chunk) == 0:
            y_cursor += padding_top
        
        current_chunk.append(cl)
        y_cursor += line_h
    
    if current_chunk:
        chunks.append(current_chunk)
    
    # Render each chunk: background first, then text
    for chunk_idx, chunk_lines in enumerate(chunks):
        if chunk_idx > 0:
            pdf.add_page()
        
        chunk_h = padding_top + len(chunk_lines) * line_h + padding_bottom
        y_start = pdf.get_y()
        
        # 1) Draw dark background
        pdf.set_fill_color(*getattr(pdf, 'color_codeblock_bg', (26, 26, 46)))
        pdf.rect(x_start, y_start, width, chunk_h, 'F')
        

        
        # 3) Render text on top of background
        pdf.set_text_color(*getattr(pdf, 'color_codeblock_text', (224, 224, 224)))
        pdf.set_font('Mono', '', 8)
        pdf.set_y(y_start + padding_top)
        
        for cl in chunk_lines:
            pdf.set_x(x_start + 8)
            pdf.cell(width - 14, line_h, cl)
            pdf.ln(line_h)
        
        pdf.set_y(y_start + chunk_h + 3)
    
    pdf.set_text_color(*getattr(pdf, 'color_text', (26, 26, 46)))
    pdf.set_font('VN', '', 11)


def _pdf_blockquote(pdf: FPDF, text: str):
    """Render a blockquote with left border."""
    pdf.ln(2)
    x_start = pdf.l_margin
    width = pdf.w - pdf.l_margin - pdf.r_margin
    max_y = pdf.h - pdf.b_margin - 18

    pdf.set_fill_color(*getattr(pdf, 'color_quote_bg', (248, 246, 255)))
    pdf.set_font('VN', 'I', 10.5)

    text_width = width - 14
    str_w = pdf.get_string_width(text)
    n_lines = max(1, int(str_w / text_width) + 1)
    line_h = 6.5
    total_h = n_lines * line_h + 8

    # Page break check
    if pdf.get_y() + total_h > max_y:
        pdf.add_page()

    y_start = pdf.get_y()
    pdf.set_fill_color(*getattr(pdf, 'color_quote_bg', (248, 246, 255)))
    pdf.rect(x_start + 5, y_start, width - 5, total_h, 'F')



    pdf.set_text_color(80, 80, 80)
    pdf.set_xy(x_start + 12, y_start + 4)
    pdf.multi_cell(text_width, line_h, text)

    pdf.set_y(y_start + total_h + 3)
    pdf.set_font('VN', '', 11)
    pdf.set_text_color(*getattr(pdf, 'color_text', (26, 26, 46)))


def _pdf_list_item(pdf: FPDF, text: str, indent: int, is_ordered: bool, marker: str):
    """Render a list item."""
    indent_x = pdf.l_margin + 6 + indent * 6
    pdf.set_x(indent_x)
    pdf.set_font('VN', '', 11)
    pdf.set_text_color(*getattr(pdf, 'color_text', (26, 26, 46)))

    bullet = marker if is_ordered else chr(8226)  # bullet character
    prefix = f"  {bullet} "
    pdf.write(5, prefix)
    _pdf_rich_text(pdf, text)


def _pdf_table(pdf: FPDF, header: list, rows: list):
    """Render a table with colored header, auto-width columns, and text wrapping."""
    pdf.ln(4)
    cols = len(header)
    if cols == 0:
        return

    avail_width = pdf.w - pdf.l_margin - pdf.r_margin
    line_h = 7
    font_size = 8.5
    max_y = pdf.h - pdf.b_margin - 18

    # --- Smart column width calculation ---
    pdf.set_font('VN', 'B', font_size)
    
    col_max_w = []
    for i in range(cols):
        hw = pdf.get_string_width(header[i]) + 6
        dw = 0
        for row in rows:
            if i < len(row):
                w = pdf.get_string_width(row[i]) + 6
                dw = max(dw, w)
        col_max_w.append(max(hw, dw))
    
    total_natural = sum(col_max_w)
    
    if total_natural <= avail_width:
        col_widths = [w * avail_width / total_natural for w in col_max_w]
    else:
        min_col = 18
        col_widths = []
        for w in col_max_w:
            ratio = w / total_natural
            cw = max(min_col, ratio * avail_width)
            col_widths.append(cw)
        s = sum(col_widths)
        col_widths = [w * avail_width / s for w in col_widths]

    def _clean_markdown(text):
        """Clean markdown markers and return plain text + whether it is bold"""
        if not text:
            return "", False
        clean = text.strip()
        is_bold = False
        if clean.startswith('**') and clean.endswith('**'):
            is_bold = True
            clean = clean[2:-2]
        else:
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', clean)
            clean = re.sub(r'\*(.*?)\*', r'\1', clean)
            clean = re.sub(r'`(.*?)`', r'\1', clean)
        return clean, is_bold

    def _draw_header():
        """Draw table header row."""
        pdf.set_font('VN', 'B', font_size)
        pdf.set_fill_color(*getattr(pdf, 'color_primary', (15, 52, 96)))
        pdf.set_text_color(255, 255, 255)
        pdf.set_draw_color(180, 180, 200)
        for i, cell_text in enumerate(header):
            pdf.cell(col_widths[i], line_h + 4, "  " + cell_text, border=1, fill=True)
        pdf.ln(line_h + 4)

    def _calc_row_height(row_data):
        """Calculate the height needed for a row with text wrapping, adding padding."""
        max_lines = 1
        for col_idx in range(cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ''
            clean_text, _ = _clean_markdown(cell_text)
            cw = col_widths[col_idx] - 4  # padding
            if cw <= 0:
                continue
            text_w = pdf.get_string_width(clean_text)
            n_lines = max(1, int(text_w / cw) + 1)
            max_lines = max(max_lines, n_lines)
        return max_lines * line_h + 4  # Thêm 4mm vertical padding

    def _draw_row(row_data, row_idx):
        """Draw a single data row with text wrapping and vertical alignment centering."""
        pdf.set_font('VN', '', font_size)
        pdf.set_text_color(*getattr(pdf, 'color_text', (26, 26, 46)))
        
        if row_idx % 2 == 1:
            pdf.set_fill_color(*getattr(pdf, 'color_table_alt', (244, 246, 250)))
        else:
            pdf.set_fill_color(255, 255, 255)
        
        row_h = _calc_row_height(row_data)
        x_start = pdf.get_x()
        y_start = pdf.get_y()
        
        # 1) Vẽ nền và border cho tất cả các ô trong hàng trước để đảm bảo thẳng hàng
        for col_idx in range(cols):
            x = x_start + sum(col_widths[:col_idx])
            pdf.set_xy(x, y_start)
            pdf.cell(col_widths[col_idx], row_h, '', border=1, fill=True)
            
        # 2) Vẽ nội dung text (multi_cell) đè lên trên các ô
        for col_idx in range(cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ''
            clean_text, is_bold = _clean_markdown(cell_text)
            
            if is_bold:
                pdf.set_font('VN', 'B', font_size)
            else:
                pdf.set_font('VN', '', font_size)
                
            x = x_start + sum(col_widths[:col_idx])
            cw = col_widths[col_idx] - 4
            text_w = pdf.get_string_width(clean_text)
            cell_lines = max(1, int(text_w / cw) + 1)
            cell_h = cell_lines * line_h
            
            # Căn giữa theo chiều dọc (Vertical Align Center)
            y_offset = (row_h - cell_h) / 2
            pdf.set_xy(x + 2, y_start + y_offset)
            pdf.multi_cell(col_widths[col_idx] - 4, line_h, clean_text)
        
        # Đưa con trỏ xuống dưới hàng vừa vẽ để tiếp tục phần sau
        pdf.set_xy(x_start, y_start + row_h)

    # --- Render table ---
    # Check if header + at least 1 row fits on current page
    header_h = line_h + 4
    first_row_h = _calc_row_height(rows[0]) if rows else 0
    if pdf.get_y() + header_h + first_row_h > max_y:
        pdf.add_page()
    
    _draw_header()
    
    pdf.set_font('VN', '', font_size)
    for row_idx, row_data in enumerate(rows):
        row_h = _calc_row_height(row_data)
        
        if pdf.get_y() + row_h > max_y:
            pdf.add_page()
            _draw_header()
        
        _draw_row(row_data, row_idx)

    pdf.ln(4)
    pdf.set_font('VN', '', 11)


# ============================================================
# DOCX CONVERSION
# ============================================================

def _setup_docx_styles(doc: Document, style_config: dict):
    """Setup custom styles for the Word document based on style metadata"""
    style = doc.styles
    font_name = style_config.get('font_family', 'Calibri')
    theme_hex = style_config.get('theme_color', '#0F3460').upper().lstrip('#')
    custom_enabled = style_config.get('custom_colors_enabled', False)
    
    # Convert hex to RGB color
    def hex_to_rgb(hex_str):
        if not hex_str:
            return RGBColor(15, 52, 96)
        hex_str = hex_str.lstrip('#')
        if len(hex_str) == 3:
            hex_str = ''.join(c*2 for c in hex_str)
        return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))
        
    if custom_enabled:
        h1_rgb = hex_to_rgb(style_config.get('h1_color', theme_hex))
        h2_rgb = hex_to_rgb(style_config.get('h2_color', theme_hex))
        h3_rgb = hex_to_rgb(style_config.get('h3_color', theme_hex))
        h4_rgb = hex_to_rgb(style_config.get('h4_color', theme_hex))
        toc_rgb = hex_to_rgb(style_config.get('toc_color', theme_hex))
    else:
        if theme_hex == 'A91D22':
            primary_rgb = RGBColor(169, 29, 34)
            h1_rgb = primary_rgb
            h2_rgb = primary_rgb
            h3_rgb = RGBColor(245, 166, 35) # #f5a623
            h4_rgb = primary_rgb
            toc_rgb = primary_rgb
        else:
            primary_rgb = hex_to_rgb(theme_hex)
            h1_rgb = primary_rgb
            h2_rgb = RGBColor(max(0, primary_rgb[0] - 20), max(0, primary_rgb[1] - 20), max(0, primary_rgb[2] - 20))
            h3_rgb = RGBColor(max(0, primary_rgb[0] - 40), max(0, primary_rgb[1] - 40), min(255, primary_rgb[2] + 40))
            h4_rgb = primary_rgb
            toc_rgb = h1_rgb

    normal = style['Normal']
    normal.font.name = font_name
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    h1 = style['Heading 1']
    h1.font.name = font_name
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = h1_rgb
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(10)

    h2 = style['Heading 2']
    h2.font.name = font_name
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = h2_rgb
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    h3 = style['Heading 3']
    h3.font.name = font_name
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = h3_rgb
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)

    h4 = style['Heading 4']
    h4.font.name = font_name
    h4.font.size = Pt(11.5)
    h4.font.bold = True
    h4.font.color.rgb = h4_rgb
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)

    # Cấu hình styles cho Table of Contents (TOC) để đồng bộ màu sắc và giãn cách dòng trong Word
    from docx.enum.style import WD_STYLE_TYPE
    
    try:
        toc_heading = style['TOC Heading']
    except KeyError:
        toc_heading = style.add_style('TOC Heading', WD_STYLE_TYPE.PARAGRAPH)
        toc_heading.base_style = style['Normal']
    
    toc_heading.font.name = font_name
    toc_heading.font.size = Pt(18)
    toc_heading.font.bold = True
    toc_heading.font.color.rgb = toc_rgb
    toc_heading.paragraph_format.space_before = Pt(12)
    toc_heading.paragraph_format.space_after = Pt(6)
    
    # Phân cấp màu sắc cho TOC
    if custom_enabled:
        toc_1_color = toc_rgb
        toc_2_color = h2_rgb
        toc_3_color = h3_rgb
    else:
        toc_1_color = toc_rgb
        toc_2_color = h2_rgb
        toc_3_color = h3_rgb

    for name, color, pt, is_bold in [('TOC 1', toc_1_color, 11, True), ('TOC 2', toc_2_color, 10, False), ('TOC 3', toc_3_color, 10, False)]:
        try:
            s = style[name]
        except KeyError:
            s = style.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            s.base_style = style['Normal']
            
        s.font.name = font_name
        s.font.size = Pt(pt)
        s.font.bold = is_bold
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(6) # Tạo khoảng giãn cách hợp lý tránh dính chữ
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.15


def _add_table_to_docx(doc: Document, header_row: list, data_rows: list, style_config: dict):
    """Add a formatted table to the document with rich text formatting (bold, italic, links)"""
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    theme_hex = style_config.get('theme_color', '0F3460').upper().lstrip('#')
    custom_enabled = style_config.get('custom_colors_enabled', False)
    
    if custom_enabled:
        table_bg_hex = style_config.get('table_header_bg', theme_hex).upper().lstrip('#')
    else:
        table_bg_hex = theme_hex
        
    if len(table_bg_hex) == 3:
        table_bg_hex = ''.join(c*2 for c in table_bg_hex)

    for i, text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        para = cell.paragraphs[0]
        para.text = "" # Clear default text
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        _add_rich_text_docx(para, text.strip(), style_config)
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): table_bg_hex,
            qn('w:val'): 'clear'
        })
        shading.append(bg)

    for row_idx, row_data in enumerate(data_rows):
        for col_idx, text in enumerate(row_data):
            if col_idx < cols:
                cell = table.rows[row_idx + 1].cells[col_idx]
                para = cell.paragraphs[0]
                para.text = "" # Clear default text
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(4)
                _add_rich_text_docx(para, text.strip(), style_config)
                for run in para.runs:
                    run.font.size = Pt(10)
                if row_idx % 2 == 1:
                    shading = cell._element.get_or_add_tcPr()
                    bg = shading.makeelement(qn('w:shd'), {
                        qn('w:fill'): 'F4F6FA',
                        qn('w:val'): 'clear'
                    })
                    shading.append(bg)

    doc.add_paragraph()


def _add_code_block(doc: Document, code_text: str, language: str = '', style_config: dict = None):
    """Add a styled code block to the document"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    pPr = para._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F0F5',
        qn('w:val'): 'clear'
    })
    pPr.append(shd)

    style_config = style_config or {}
    theme_hex = style_config.get('theme_color', 'E94560').upper().lstrip('#')
    custom_enabled = style_config.get('custom_colors_enabled', False)
    
    if custom_enabled:
        color_hex = style_config.get('hr_color', theme_hex).upper().lstrip('#')
    else:
        color_hex = 'F5A623' if theme_hex.upper().lstrip('#') == 'A91D22' else theme_hex

    if len(color_hex) == 3:
        color_hex = ''.join(c*2 for c in color_hex)

    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    left = pBdr.makeelement(qn('w:left'), {
        qn('w:val'): 'single',
        qn('w:sz'): '18',
        qn('w:space'): '4',
        qn('w:color'): color_hex
    })
    pBdr.append(left)
    pPr.append(pBdr)

    run = para.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)


def _setup_docx_header_footer(doc: Document, metadata: dict):
    project_name = metadata.get('project_name', '')
    version = metadata.get('doc_version', '')
    logo_path = metadata.get('logo_path', '')  # Logo bên phải
    logo_left_path = metadata.get('logo_left_path', '')  # Logo bên trái
    style = metadata.get('style', {})
    footer_format = style.get('footer_format', 'page_of_total')
    
    # Kích hoạt tự động update các trường (TOC, số trang) khi mở file Word
    try:
        settings = doc.settings._element
        updateFields = settings.makeelement(qn('w:updateFields'), {qn('w:val'): 'true'})
        settings.append(updateFields)
    except Exception:
        pass
        
    def add_page_number_to_run(run):
        fldSimple = run._element.makeelement(qn('w:fldSimple'), {qn('w:instr'): 'PAGE'})
        run._element.append(fldSimple)

    def add_numpages_to_run(run):
        fldSimple = run._element.makeelement(qn('w:fldSimple'), {qn('w:instr'): 'NUMPAGES'})
        run._element.append(fldSimple)
    
    for section in doc.sections:
        # Header - Sử dụng Table 3 cột không viền để tránh lỗi căn giữa tab stop
        header = section.header
        
        # Clear default paragraph
        header_para = header.paragraphs[0]
        header_para.text = ""
        
        # Tính toán độ rộng khả dụng
        avail_width = section.page_width - section.left_margin - section.right_margin
        
        # Tạo bảng 3 cột cho Header
        header_table = header.add_table(rows=1, cols=3, width=avail_width)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_table.style = 'Normal Table'  # Bảng không viền mặc định
        
        # Set độ rộng cột
        col_l_w = Cm(1.5)
        col_r_w = Cm(1.5)
        col_m_w = avail_width - col_l_w - col_r_w
        
        header_table.columns[0].width = col_l_w
        header_table.columns[1].width = col_m_w
        header_table.columns[2].width = col_r_w
        
        for i, w in enumerate([col_l_w, col_m_w, col_r_w]):
            header_table.rows[0].cells[i].width = w
            
        cell_l = header_table.rows[0].cells[0]
        cell_m = header_table.rows[0].cells[1]
        cell_r = header_table.rows[0].cells[2]
        
        p_l = cell_l.paragraphs[0]
        p_l.text = ""
        p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        p_m = cell_m.paragraphs[0]
        p_m.text = ""
        p_m.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        p_r = cell_r.paragraphs[0]
        p_r.text = ""
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Chèn Logo Trái
        if logo_left_path and os.path.exists(logo_left_path):
            try:
                run_logo_l = p_l.add_run()
                run_logo_l.add_picture(logo_left_path, height=Cm(0.6))
            except Exception:
                pass
                
        # Chèn Tiêu đề dự án
        header_text = project_name
        if version:
            header_text += f" - Phiên bản: {version}"
            
        if header_text:
            run_m = p_m.add_run(header_text)
            run_m.font.size = Pt(8.5)
            run_m.font.color.rgb = RGBColor(120, 120, 120)
            run_m.italic = True
            
        # Chèn Logo Phải
        if logo_path and os.path.exists(logo_path):
            try:
                run_logo_r = p_r.add_run()
                run_logo_r.add_picture(logo_path, height=Cm(0.6))
            except Exception:
                pass
                
        # Thêm dòng kẻ dưới Header tinh tế
        p_line = header.add_paragraph()
        p_line.paragraph_format.space_before = Pt(2)
        p_line.paragraph_format.space_after = Pt(0)
        p_line.paragraph_format.line_spacing = Pt(1)
        run_line = p_line.add_run()
        run_line.font.size = Pt(1)
        
        pPr = p_line._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',      # 0.5 pt
            qn('w:space'): '1',
            qn('w:color'): 'D3D3D3'
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
            
        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "" # Clear default
        
        if footer_format == 'right_align':
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Chèn số trang theo định dạng cấu hình
        run_f = footer_para.add_run()
        run_f.font.size = Pt(8.5)
        run_f.font.color.rgb = RGBColor(120, 120, 120)
        run_f.italic = True
        
        if footer_format == 'page_only':
            add_page_number_to_run(run_f)
        elif footer_format == 'brackets':
            run_f.text = "- "
            add_page_number_to_run(footer_para.add_run())
            run_end = footer_para.add_run(" -")
            run_end.font.size = Pt(8.5)
            run_end.font.color.rgb = RGBColor(120, 120, 120)
            run_end.italic = True
        else: # page_of_total hoặc right_align
            run_f.text = "Trang "
            add_page_number_to_run(footer_para.add_run())
            
            run_mid = footer_para.add_run("/")
            run_mid.font.size = Pt(8.5)
            run_mid.font.color.rgb = RGBColor(120, 120, 120)
            run_mid.italic = True
            
            add_numpages_to_run(footer_para.add_run())


def _add_toc_to_docx(doc: Document, style_config: dict):
    """Add a self-updating Table of Contents field at the beginning of the DOCX"""
    para = doc.add_paragraph()
    run = para.add_run("Mục lục")
    run.font.size = Pt(18)
    run.bold = True
    
    theme_hex = style_config.get('theme_color', '0F3460').upper().lstrip('#')
    custom_enabled = style_config.get('custom_colors_enabled', False)
    
    if custom_enabled:
        toc_hex = style_config.get('toc_color', theme_hex).upper().lstrip('#')
        border_hex = style_config.get('hr_color', theme_hex).upper().lstrip('#')
    else:
        toc_hex = 'A91D22' if theme_hex.upper().lstrip('#') == 'A91D22' else theme_hex
        border_hex = 'F5A623' if theme_hex.upper().lstrip('#') == 'A91D22' else theme_hex
        
    if len(toc_hex) == 3:
        toc_hex = ''.join(c*2 for c in toc_hex)
    if len(border_hex) == 3:
        border_hex = ''.join(c*2 for c in border_hex)
        
    def hex_to_rgb(hex_str):
        return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))
        
    run.font.color.rgb = hex_to_rgb(toc_hex)
    
    # Add horizontal line
    para_border = doc.add_paragraph()
    pPr = para_border._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '12',
        qn('w:space'): '1',
        qn('w:color'): border_hex
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    # Insert TOC field XML theo thứ tự tuyến tính chuẩn xác
    toc_para = doc.add_paragraph()
    
    run_begin = toc_para.add_run()
    fldChar1 = run_begin._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run_begin._element.append(fldChar1)
    
    run_instr = toc_para.add_run()
    instrText = run_instr._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run_instr._element.append(instrText)
    
    run_sep = toc_para.add_run()
    fldChar2 = run_sep._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    run_sep._element.append(fldChar2)
    
    placeholder = toc_para.add_run(" (Nhấp chuột phải và chọn 'Update Field' để cập nhật mục lục) ")
    placeholder.italic = True
    placeholder.font.color.rgb = RGBColor(120, 120, 120)
    
    run_end = toc_para.add_run()
    fldChar3 = run_end._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run_end._element.append(fldChar3)
    
    doc.add_page_break()


def _docx_image(doc: Document, img_path: str, alt_text: str):
    """Render image in Word document, scaling to fit page width safely"""
    local_path = _get_image_file(img_path)
    if not local_path or not os.path.exists(local_path):
        para = doc.add_paragraph()
        run = para.add_run(f"[Hình ảnh không tải được: {alt_text or img_path}]")
        run.italic = True
        run.font.color.rgb = RGBColor(200, 50, 50)
        return
        
    try:
        doc.add_picture(local_path, width=Cm(15))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"Error rendering image in Word: {e}")
        para = doc.add_paragraph()
        run = para.add_run(f"[Lỗi hiển thị hình ảnh: {alt_text or img_path}]")
        run.italic = True
        run.font.color.rgb = RGBColor(200, 50, 50)


def convert_md_to_docx(md_content: str, title: str = "Document", metadata: dict = None) -> bytes:
    """Convert Markdown content to DOCX bytes"""
    doc = Document()
    metadata = metadata or {}
    style = metadata.get('style', {})
    
    # Định dạng lề
    margin_type = style.get('margin', 'standard')
    margin_val_cm = {
        'narrow': 1.5,
        'standard': 2.0,
        'wide': 2.54
    }.get(margin_type, 2.0)
    
    # Hướng giấy & Khổ giấy
    page_size = style.get('page_size', 'A4')
    orientation = style.get('orientation', 'P')
    
    for section in doc.sections:
        # Lề
        section.top_margin = Cm(margin_val_cm)
        section.bottom_margin = Cm(margin_val_cm)
        section.left_margin = Cm(margin_val_cm)
        section.right_margin = Cm(margin_val_cm)
        
        # Hướng giấy & Cỡ giấy
        if page_size == 'Letter':
            w_cm, h_cm = 21.59, 27.94
        else: # A4
            w_cm, h_cm = 21.0, 29.7
            
        if orientation == 'L':
            section.page_width = Cm(h_cm)
            section.page_height = Cm(w_cm)
            section.orientation = 1 # landscape
        else:
            section.page_width = Cm(w_cm)
            section.page_height = Cm(h_cm)
            section.orientation = 0 # portrait

    _setup_docx_styles(doc, style)
    _setup_docx_header_footer(doc, metadata)

    lines = md_content.split('\n')
    
    # Check if there are headings to insert TOC
    has_headings = any(re.match(r'^(#{1,3})\s+(.*)', line.strip()) for line in lines)
    if has_headings:
        _add_toc_to_docx(doc, style)

    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ''
    in_table = False
    table_header = []
    table_rows = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith('```'):
            if in_code_block:
                _add_code_block(doc, '\n'.join(code_lines), code_lang, style)
                code_lines = []
                code_lang = ''
                in_code_block = False
            else:
                if in_table and table_header:
                    _add_table_to_docx(doc, table_header, table_rows, style)
                    table_header, table_rows = [], []
                    in_table = False
                in_code_block = True
                lang_match = re.match(r'^```(\w+)', line.strip())
                code_lang = lang_match.group(1) if lang_match else ''
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_header = cells
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table and table_header:
                _add_table_to_docx(doc, table_header, table_rows, style)
                table_header, table_rows = [], []
                in_table = False

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if re.match(r'^[-*_]{3,}\s*$', stripped):
            para = doc.add_paragraph()
            pPr = para._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            theme_hex = style.get('theme_color', 'E94560').upper().lstrip('#')
            custom_enabled = style.get('custom_colors_enabled', False)
            if custom_enabled:
                color_hex = style.get('hr_color', theme_hex).upper().lstrip('#')
            else:
                color_hex = 'F5A623' if theme_hex.upper().lstrip('#') == 'A91D22' else theme_hex
                
            if len(color_hex) == 3:
                color_hex = ''.join(c*2 for c in color_hex)
                
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '12',
                qn('w:space'): '1',
                qn('w:color'): color_hex
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Image block
        image_match = re.match(r'^!\[(.*?)\]\((.*?)\)', stripped)
        if image_match:
            alt = image_match.group(1)
            img_path = image_match.group(2)
            _docx_image(doc, img_path, alt)
            i += 1
            continue

        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            if level <= 4:
                doc.add_heading(text, level=level)
            else:
                para = doc.add_paragraph(text)
                para.runs[0].bold = True
                para.runs[0].font.size = Pt(11)
            i += 1
            continue

        if stripped.startswith('>'):
            quote_text = re.sub(r'^>\s*', '', stripped)
            quote_text = re.sub(r'\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', r'[\1]', quote_text)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            pPr = para._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            
            theme_hex = style.get('theme_color', '533483').upper().lstrip('#')
            custom_enabled = style.get('custom_colors_enabled', False)
            if custom_enabled:
                color_hex = style.get('quote_border', theme_hex).upper().lstrip('#')
                bg_hex = style.get('quote_bg', 'F8F6FF').upper().lstrip('#')
            else:
                color_hex = 'F5A623' if theme_hex.upper().lstrip('#') == 'A91D22' else '533483'
                bg_hex = 'FEFBEC' if theme_hex.upper().lstrip('#') == 'A91D22' else 'F8F6FF'
                
            if len(color_hex) == 3:
                color_hex = ''.join(c*2 for c in color_hex)
            if len(bg_hex) == 3:
                bg_hex = ''.join(c*2 for c in bg_hex)
                
            # Shading for blockquote
            shd = pPr.makeelement(qn('w:shd'), {
                qn('w:fill'): bg_hex,
                qn('w:val'): 'clear'
            })
            pPr.append(shd)
            
            left = pBdr.makeelement(qn('w:left'), {
                qn('w:val'): 'single',
                qn('w:sz'): '24',
                qn('w:space'): '4',
                qn('w:color'): color_hex
            })
            pBdr.append(left)
            pPr.append(pBdr)
            
            run = para.add_run(quote_text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            i += 1
            continue

        list_match = re.match(r'^(\s*)([-*+]|\d+[.)])\s+(.*)', stripped)
        if list_match:
            indent = len(list_match.group(1))
            marker = list_match.group(2)
            text = list_match.group(3)
            is_ordered = bool(re.match(r'\d+[.)]', marker))
            style_name = 'List Number' if is_ordered else 'List Bullet'
            para = doc.add_paragraph(style=style_name)
            _add_rich_text_docx(para, text, style)
            if indent > 0:
                para.paragraph_format.left_indent = Cm(1 + indent * 0.5)
            i += 1
            continue

        para = doc.add_paragraph()
        _add_rich_text_docx(para, stripped, style)
        i += 1

    if in_table and table_header:
        _add_table_to_docx(doc, table_header, table_rows, style)
    if in_code_block and code_lines:
        _add_code_block(doc, '\n'.join(code_lines), code_lang, style)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _add_rich_text_docx(para, text: str, style_config: dict = None):
    """Add text with inline markdown formatting (bold, italic, code, links)"""
    style_config = style_config or {}
    custom_enabled = style_config.get('custom_colors_enabled', False)
    theme_hex = style_config.get('theme_color', '#0f3460').upper().lstrip('#')
    
    def hex_to_rgb(hex_str):
        if not hex_str:
            return RGBColor(15, 52, 96)
        hex_str = hex_str.lstrip('#')
        if len(hex_str) == 3:
            hex_str = ''.join(c*2 for c in hex_str)
        return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))
        
    if custom_enabled:
        bold_rgb = hex_to_rgb(style_config.get('bold_color', theme_hex))
        link_rgb = hex_to_rgb(style_config.get('h1_color', theme_hex)) # default link color to H1
    else:
        if theme_hex == 'A91D22':
            bold_rgb = RGBColor(169, 29, 34)
            link_rgb = RGBColor(169, 29, 34)
        else:
            primary_rgb = hex_to_rgb(theme_hex)
            bold_rgb = RGBColor(max(0, primary_rgb[0] - 20), max(0, primary_rgb[1] - 20), max(0, primary_rgb[2] - 20))
            link_rgb = primary_rgb

    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'
        r'|(\*\*(.+?)\*\*)'
        r'|(\*(.+?)\*)'
        r'|(`(.+?)`)'
        r'|(\[(.+?)\]\((.+?)\))'
    )

    last_end = 0
    for match in pattern.finditer(text):
        if match.start() > last_end:
            para.add_run(text[last_end:match.start()])

        if match.group(2):
            run = para.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(4):
            run = para.add_run(match.group(4))
            run.bold = True
            run.font.color.rgb = bold_rgb
        elif match.group(6):
            run = para.add_run(match.group(6))
            run.italic = True
            run.font.color.rgb = RGBColor(0x53, 0x34, 0x83)
        elif match.group(8):
            run = para.add_run(match.group(8))
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xe9, 0x45, 0x60)
        elif match.group(10):
            link_text = match.group(10)
            run = para.add_run(link_text)
            run.font.color.rgb = link_rgb
            run.underline = True

        last_end = match.end()

    if last_end < len(text):
        para.add_run(text[last_end:])
